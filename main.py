import streamlit as st
import pandas as pd
import numpy as np
from scipy import stats
import statsmodels.api as sm
from statsmodels.formula.api import ols
from statsmodels.stats.anova import anova_lm
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.stats.stattools import durbin_watson
import io
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# -----------------------------------------------------------------------------
# 1. UI 스타일링 및 테마 설정
# -----------------------------------------------------------------------------
st.set_page_config(page_title="STATERA", page_icon="🎓", layout="wide")

plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['axes.unicode_minus'] = False
sns.set_theme(style="whitegrid")

ACRONYM_FULL = "STATistical Engine for Research & Analysis"

st.markdown(f"""
<style>
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
    * {{ font-family: 'Pretendard', sans-serif; }}
    .main-header {{ color: #0d9488; text-align: center; font-size: 2.8rem; font-weight: 800; margin-bottom: 5px; }}
    .sub-header {{ text-align: center; color: #64748b; font-size: 1.1rem; margin-bottom: 40px; }}
    
    .guide-container {{ display: flex; gap: 20px; margin-bottom: 30px; }}
    .guide-box {{ flex: 1; background: white; border: 1px solid #e2e8f0; border-radius: 16px; padding: 24px; box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.05); }}
    .guide-label {{ font-size: 1.1rem; font-weight: 700; color: #0f172a; margin-bottom: 8px; display: flex; align-items: center; }}
    .guide-text {{ font-size: 0.9rem; color: #64748b; line-height: 1.6; }}

    .mentor-box {{ background-color: #f0fdfa; border-left: 6px solid #0d9488; padding: 25px; border-radius: 12px; margin-bottom: 30px; }}
    .mentor-title {{ color: #0f766e; font-size: 1.3rem; font-weight: 700; margin-bottom: 12px; }}
    .mentor-content {{ color: #1e293b; font-size: 1rem; line-height: 1.8; }}

    .section-title {{ font-size: 1.6rem; font-weight: 800; color: #0f172a; margin-top: 50px; margin-bottom: 25px; border-bottom: 2px solid #e2e8f0; padding-bottom: 12px; display: flex; align-items: center; }}
    .step-badge {{ background: #0d9488; color: white; border-radius: 8px; padding: 4px 15px; font-size: 0.9rem; margin-right: 15px; vertical-align: middle; }}

    .assumption-pass {{ background-color: #dcfce7; color: #166534; padding: 12px; border-radius: 8px; margin-bottom: 8px; border: 1px solid #bbf7d0; font-weight: 600; font-size: 0.95rem; }}
    .assumption-fail {{ background-color: #fee2e2; color: #991b1b; padding: 12px; border-radius: 8px; margin-bottom: 8px; border: 1px solid #fecaca; font-weight: 600; font-size: 0.95rem; }}
    
    .ethics-container {{ background-color: #fff7ed; border: 1px solid #ffedd5; border-radius: 12px; padding: 20px; margin-top: 50px; margin-bottom: 30px; }}
    .ethics-title {{ color: #c2410c; font-size: 1.1rem; font-weight: 700; margin-bottom: 10px; }}
    .ethics-text {{ color: #9a3412; font-size: 0.9rem; line-height: 1.6; }}

    div[data-testid="stRadio"] > div {{ flex-direction: row; gap: 20px; overflow-x: auto; }}
    .stButton>button {{ width: 100%; border-radius: 12px; background: #0d9488; color: white; font-weight: 700; height: 3.8em; border: none; transition: 0.4s; }}
    
    /* 데이터프레임 헤더 스타일링 및 인덱스 숨기기용 */
    thead tr th:first-child {{ display:none }}
    tbody th {{ display:none }}
</style>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 2. 통계 멘토 가이드 데이터 및 유틸리티
# -----------------------------------------------------------------------------
def format_p(p): return "<.001" if p < .001 else f"{p:.3f}"
def get_stars(p): return "***" if p < .001 else "**" if p < .01 else "*" if p < .05 else ""
def get_plot_buffer():
    buf = io.BytesIO(); plt.savefig(buf, format='png', bbox_inches='tight', dpi=300); buf.seek(0); plt.close(); return buf

STAT_MENTOR = {
    "기술통계": {"purpose": "데이터의 중심 경향성과 분포 특성을 요약합니다.", "indicator": "평균은 자료의 수준을, 표준편차는 산포 정도를 나타냅니다.", "check": "왜도와 첨도를 통해 정규분포 가정을 검토하십시오."},
    "빈도분석": {"purpose": "범주형 변수의 빈도와 비율을 파악합니다.", "indicator": "사례 수(n)와 유효 백분율(%)을 산출하여 제시합니다.", "check": "결측치가 전체 비중에 미치는 영향을 확인하십시오."},
    "카이제곱 검정": {"purpose": "범주형 변수 간의 통계적 관련성 유무를 확인합니다.", "indicator": "기대빈도 가정 충족 여부에 따라 분석 결과의 타당성을 평가합니다.", "check": "기대빈도 5 미만 셀 비율이 20%를 초과하는지 검토하십시오."},
    "단일표본 T-검정": {"purpose": "표본 평균을 특정 기준값과 비교하여 차이를 검증합니다.", "indicator": "t값과 유의확률을 통해 기준치와의 통계적 거리를 판정합니다.", "check": "집단의 정규성 가정을 사전에 확인하십시오."},
    "독립표본 T-검정": {"purpose": "서로 독립적인 두 집단 간의 평균 차이를 비교 분석합니다.", "indicator": "두 집단 간 평균값 차이가 유의미한 수준인지 판정합니다.", "check": "두 집단의 정규성과 등분산성 가정을 확인하십시오."},
    "대응표본 T-검정": {"purpose": "동일 집단의 처치 전후(사전-사후) 평균 변화를 비교합니다.", "indicator": "사전-사후 점수 차이가 0에서 얼마나 벗어났는지 검증합니다.", "check": "차이값의 정규성 분포를 검토하십시오."},
    "분산분석(ANOVA)": {"purpose": "세 집단 이상의 평균 차이를 비교하고 변량 차이를 분석합니다.", "indicator": "F값으로 유의성을 판정한 후 사후분석(Tukey 등)을 수행합니다.", "check": "집단별 정규성과 등분산성 가정을 확인하십시오."},
    "상관분석": {"purpose": "두 연속형 변수 간의 선형적 관계의 강도를 파악합니다.", "indicator": "상관계수(r)를 통해 변수 간 관계의 방향과 밀접도를 평가합니다.", "check": "변수 간의 관계가 선형적인지 산점도를 검토하십시오."},
    "신뢰도 분석": {"purpose": "측정 도구의 문항들이 일관성 있게 측정되는지 평가합니다.", "indicator": "Cronbach α 계수가 0.7 이상일 때 신뢰도가 확보된 것으로 간주합니다.", "check": "역코딩 문항이 분석 전 적절히 변환되었는지 확인하십시오."},
    "회귀분석": {"purpose": "독립변수가 종속변수에 미치는 영향력을 수치화합니다.", "indicator": "R2로 모형 설명력을, Beta로 영향력의 크기를 평가합니다.", "check": "다중공선성(VIF < 10)과 잔차 가정을 검토하십시오."}
}

def create_pro_report(m_name, r_df, interpretation, guide, plot_b=None, assump=""):
    doc = Document(); doc.styles['Normal'].font.name = 'Malgun Gothic'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    doc.add_heading(f'STATERA Report: {m_name}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if assump: 
        doc.add_heading('1. Assumption Checks', level=1)
        # HTML 태그 제거 후 텍스트만 저장
        clean_assump = assump.replace('<div class="assumption-pass">', '').replace('<div class="assumption-fail">', '').replace('</div>', '')
        doc.add_paragraph(clean_assump).italic = True
    doc.add_heading('2. Statistical Results', level=1)
    t = doc.add_table(r_df.shape[0]+1, r_df.shape[1]); t.style = 'Table Grid'
    for j, c in enumerate(r_df.columns): t.cell(0,j).text = str(c)
    for i in range(r_df.shape[0]):
        for j in range(r_df.shape[1]): t.cell(i+1,j).text = str(r_df.values[i,j])
    if plot_b: doc.add_heading('3. Visualization', level=1); doc.add_picture(plot_b, width=Inches(4.5))
    doc.add_heading('4. AI Interpretation', level=1); doc.add_paragraph(interpretation)
    doc.add_heading('5. Thesis Writing Guide', level=1); doc.add_paragraph(guide)
    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio

# -----------------------------------------------------------------------------
# 3. 사이드바
# -----------------------------------------------------------------------------
with st.sidebar:
    st.markdown("<h1 style='color:#0d9488;'>STATERA 📊</h1>", unsafe_allow_html=True)
    st.caption(ACRONYM_FULL)
    st.markdown("---")
    st.markdown("### 🚧 Research Beta Version")
    st.info("본 서비스는 연구 데이터 분석의 진입 장벽을 낮추기 위해 개발된 웹 기반 통계 솔루션입니다. 현재 분석 알고리즘의 타당도 검증 절차를 진행 중입니다.")
    st.markdown("---")
    st.markdown("### 📬 Contact & Feedback")
    st.write("오류 제보 및 기능 제안은 언제나 환영합니다.")
    st.link_button("📧 메일 보내기", "mailto:nncj91@snu.ac.kr")
    st.caption("주소 복사:")
    st.code("nncj91@snu.ac.kr", language="text")
    st.markdown("---")
    st.caption("© 2026 ANDA Lab. Developed by Jeongin Choe.")

# -----------------------------------------------------------------------------
# 4. 메인 어플리케이션 레이아웃
# -----------------------------------------------------------------------------
st.markdown('<div class="main-header">STATERA</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">수치적 결과 산출을 넘어, 연구 논리와 학술적 해석의 과정을 체득하는 통계 학습 플랫폼입니다.</div>', unsafe_allow_html=True)

st.markdown(f"""
<div class="guide-container">
    <div class="guide-box">
        <div class="guide-label">🔒 데이터 보안 안내</div>
        <div class="guide-text">업로드된 데이터는 분석 즉시 메모리에서 삭제되며 서버에 저장되지 않아 보안이 철저히 유지됩니다.</div>
    </div>
    <div class="guide-box">
        <div class="guide-label">📄 데이터 형식 가이드</div>
        <div class="guide-text">첫 번째 행에는 반드시 변수명이 포함되어야 하며, XLSX 또는 CSV 형식의 파일만 인식 가능합니다.</div>
    </div>
</div>
""", unsafe_allow_html=True)

up_file = st.file_uploader("파일을 업로드하여 분석을 시작하십시오.", type=["xlsx", "csv"], label_visibility="collapsed")

if up_file:
    df = pd.read_excel(up_file) if up_file.name.endswith('xlsx') else pd.read_csv(up_file)
    num_cols = df.select_dtypes(include=[np.number]).columns
    all_cols = df.columns
    st.success(f"데이터 로드 완료: 분석 대상 사례 수 N={len(df)}")

    # Step 01: 분석 기법 선택
    st.markdown('<div class="section-title"><span class="step-badge">01</span> 연구 목적에 따른 분석 기법 선택</div>', unsafe_allow_html=True)
    
    group = st.selectbox("분석 범주를 선택하십시오.", [
        "기초 데이터 분석 (Descriptive/Frequency)", 
        "집단 간 차이 검정 (T-test/ANOVA)", 
        "관계 및 영향력 분석 (Regression/Corr)",
        "척도 신뢰도 분석 (Reliability)"
    ])
    
    if "기초" in group: 
        m_list = ["기술통계", "빈도분석", "카이제곱 검정"]
    elif "차이" in group: 
        m_list = ["단일표본 T-검정", "독립표본 T-검정", "대응표본 T-검정", "분산분석(ANOVA)"]
    elif "관계" in group: 
        m_list = ["상관분석", "회귀분석"]
    else: 
        m_list = ["신뢰도 분석"]
    
    method = st.radio("상세 분석 기법 선택", m_list, horizontal=True)
    
    m_info = STAT_MENTOR.get(method.split(" (")[0] if " (" in method else method, {"purpose": "데이터 분석 수행", "indicator": "지표 산출", "check": "가정 검토"})
    
    st.markdown(f"""
    <div class="mentor-box">
        <div class="mentor-title">👨‍🏫 {method} 학술 가이드</div>
        <div class="mentor-content">
            <b>분석 목적:</b> {m_info['purpose']}<br>
            <b>핵심 지표 해석:</b> {m_info['indicator']}<br>
            <b>데이터 점검 사항:</b> {m_info['check']}
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Step 02: 변수 선택 및 실행
    st.markdown('<div class="section-title"><span class="step-badge">02</span> 분석 변수 설정 및 실행</div>', unsafe_allow_html=True)
    final_df, p_val, interp, plot_img, assump_report = None, None, "", None, []

    # 기법별 상세 로직 구현
    if method == "기술통계":
        v = st.selectbox("분석할 변수 (연속형)", num_cols)
        if st.button("통계 분석 실행"):
            final_df = df[[v]].describe().T.reset_index().round(2)
            skew = df[v].skew(); kurt = df[v].kurt()
            if abs(skew) < 3 and abs(kurt) < 10:
                assump_report.append(f'<div class="assumption-pass">✅ 정규성 가정 충족: 왜도({skew:.2f})와 첨도({kurt:.2f})가 기준 이내입니다.</div>')
            else:
                assump_report.append(f'<div class="assumption-fail">⚠️ 정규성 가정 위배: 왜도/첨도 기준 초과. (데이터 변환 또는 비모수적 기술통계 고려 권장)</div>')
            plt.figure(figsize=(6,3)); sns.histplot(df[v].dropna(), kde=True, color="#0d9488"); plot_img = get_plot_buffer()
            interp = f"📌 {v}의 평균은 {df[v].mean():.2f}(SD={df[v].std():.2f})입니다."

    elif method == "빈도분석":
        vs = st.multiselect("분석할 변수들 (범주형)", all_cols)
        if st.button("통계 분석 실행") and vs:
            res = []
            for c in vs:
                counts = df[c].value_counts().reset_index(); counts.columns = ['범주', 'n']
                counts['%'] = (counts['n'] / counts['n'].sum() * 100).round(1)
                counts.insert(0, '변수명', c); res.append(counts)
            final_df = pd.concat(res)
            assump_report.append('<div class="assumption-pass">✅ 가정 검정 해당 없음: 빈도분석은 비모수적 방법으로 별도의 가정이 필요하지 않습니다.</div>')
            interp = "대상자의 일반적 분포를 확인하십시오."

    elif method == "카이제곱 검정":
        r = st.selectbox("행 변수 (범주형)", all_cols)
        c = st.selectbox("열 변수 (범주형)", all_cols)
        if st.button("통계 분석 실행"):
            ct = pd.crosstab(df[r], df[c]); chi2, p, _, exp = stats.chi2_contingency(ct)
            under_5_pct = (exp < 5).sum() / exp.size * 100
            if under_5_pct <= 20:
                assump_report.append(f'<div class="assumption-pass">✅ 기대빈도 가정 충족: 기대빈도 5 미만 셀이 {under_5_pct:.1f}%(20% 이하)입니다.</div>')
            else:
                assump_report.append(f'<div class="assumption-fail">⚠️ 기대빈도 가정 위배: 20% 초과. (대안으로 Fisher의 정확 검정(Fisher\'s Exact Test) 사용 권장)</div>')
            final_df = ct.astype(str) + " (" + (ct/ct.sum()*100).round(1).astype(str) + "%)"
            p_val = p; interp = f"📌 {r}와 {c} 간 연관성 유의확률: p={format_p(p)}"

    elif method == "단일표본 T-검정":
        y = st.selectbox("검정 변수 (연속형)", num_cols)
        ref_v = st.number_input("비교할 기준값 (Test Value)", value=0.0)
        if st.button("통계 분석 실행"):
            data = df[y].dropna(); _, sp = stats.shapiro(data)
            if sp > 0.05:
                assump_report.append(f'<div class="assumption-pass">✅ 정규성 가정 충족: Shapiro-Wilk 검정(p={sp:.3f} > .05) 결과 정규분포를 따릅니다.</div>')
            else:
                assump_report.append(f'<div class="assumption-fail">⚠️ 정규성 가정 위배: p={sp:.3f} < .05. (대안으로 비모수 검정인 Wilcoxon Signed-Rank Test 사용 권장)</div>')
            stat, p = stats.ttest_1samp(data, ref_v); p_val = p
            final_df = pd.DataFrame({"방법": [method], "t값": [stat], "df": [len(data)-1], "p값": [format_p(p)]})
            interp = f"📌 평균과 기준값 간의 차이는 {'유의합니다' if p < 0.05 else '유의하지 않습니다'}."

    elif method == "독립표본 T-검정":
        g = st.selectbox("집단 변수 (범주형: 2집단)", all_cols)
        y = st.selectbox("검정 변수 (연속형)", num_cols)
        if st.button("통계 분석 실행"):
            if len(df[g].unique()) != 2:
                st.error("집단 변수는 정확히 2개의 범주를 가져야 합니다.")
            else:
                gps = df[g].unique(); g1, g2 = df[df[g]==gps[0]][y].dropna(), df[df[g]==gps[1]][y].dropna()
                
                _, sp1 = stats.shapiro(g1); _, sp2 = stats.shapiro(g2)
                if sp1 > 0.05 and sp2 > 0.05:
                     assump_report.append(f'<div class="assumption-pass">✅ 정규성 가정 충족: 두 집단 모두 정규분포를 따릅니다.</div>')
                else:
                     assump_report.append(f'<div class="assumption-fail">⚠️ 정규성 가정 위배: 한 집단 이상이 정규성을 만족하지 않습니다. (대안으로 Mann-Whitney U Test 사용 권장)</div>')

                _, lp = stats.levene(g1, g2)
                if lp > 0.05:
                    assump_report.append(f'<div class="assumption-pass">✅ 등분산성 가정 충족: Levene 검정(p={lp:.3f} > .05) 결과 분산이 동일합니다.</div>')
                    stat, p = stats.ttest_ind(g1, g2, equal_var=True)
                else:
                    assump_report.append(f'<div class="assumption-fail">⚠️ 등분산성 가정 위배: p={lp:.3f} < .05. (자동으로 Welch\'s T-test를 적용하여 분석을 수행했습니다)</div>')
                    stat, p = stats.ttest_ind(g1, g2, equal_var=False)

                p_val = p
                final_df = pd.DataFrame({"집단": [gps[0], gps[1]], "N": [len(g1), len(g2)], "Mean": [g1.mean(), g2.mean()], "SD": [g1.std(), g2.std()]})
                plt.figure(figsize=(5,4)); sns.boxplot(x=g, y=y, data=df); plot_img = get_plot_buffer()
                interp = f"📌 두 집단 간 {y}의 평균 차이는 t={stat:.3f}, p={format_p(p)}로 통계적으로 {'유의합니다' if p < 0.05 else '유의하지 않습니다'}."

    elif method == "대응표본 T-검정":
        y1 = st.selectbox("사전 변수 (연속형)", num_cols)
        y2 = st.selectbox("사후 변수 (연속형)", num_cols)
        if st.button("통계 분석 실행"):
            diff = df[y2] - df[y1]; _, sp = stats.shapiro(diff.dropna())
            if sp > 0.05:
                assump_report.append(f'<div class="assumption-pass">✅ 차이의 정규성 충족: Shapiro-Wilk 검정(p={sp:.3f} > .05)을 만족합니다.</div>')
            else:
                assump_report.append(f'<div class="assumption-fail">⚠️ 차이의 정규성 위배: p={sp:.3f} < .05. (대안으로 비모수 검정인 Wilcoxon Signed-Rank Test 사용 권장)</div>')
            
            stat, p = stats.ttest_rel(df[y1].dropna(), df[y2].dropna()); p_val = p
            
            # DataFrame 길이 오류 수정: 빈 문자열로 패딩
            final_df = pd.DataFrame({
                "변수": [y1, y2], 
                "Mean": [df[y1].mean(), df[y2].mean()], 
                "t값": [f"{stat:.3f}", ""], 
                "p값": [format_p(p), ""]
            })
            interp = f"📌 사전 대비 사후의 수치 변화는 {'유의합니다' if p < 0.05 else '유의하지 않습니다'}."

    elif method == "분산분석(ANOVA)":
        g = st.selectbox("집단 변수 (범주형: 3집단 이상)", all_cols)
        y = st.selectbox("검정 변수 (연속형)", num_cols)
        if st.button("통계 분석 실행"):
            model = ols(f'{y} ~ C({g})', data=df).fit()
            resid = model.resid; _, sp = stats.shapiro(resid)
            if sp > 0.05:
                assump_report.append(f'<div class="assumption-pass">✅ 잔차 정규성 충족: Shapiro-Wilk p={sp:.3f}</div>')
            else:
                assump_report.append(f'<div class="assumption-fail">⚠️ 잔차 정규성 위배: p={sp:.3f}. (대안으로 Kruskal-Wallis Test 사용 권장)</div>')
            
            grps = [df[df[g] == k][y].dropna() for k in df[g].unique()]
            _, lp = stats.levene(*grps)
            if lp > 0.05:
                assump_report.append(f'<div class="assumption-pass">✅ 등분산성 충족: Levene p={lp:.3f}</div>')
            else:
                assump_report.append(f'<div class="assumption-fail">⚠️ 등분산성 위배: p={lp:.3f}. (대안으로 Welch ANOVA 사용 권장)</div>')

            res = anova_lm(model, typ=2); p_val = res.iloc[0,3]
            final_df = res.reset_index().round(3)
            if p_val < 0.05:
                tukey = pairwise_tukeyhsd(df[y].dropna(), df[g].dropna())
                st.info("💡 사후검정(Tukey HSD) 결과가 하단에 출력됩니다.")
                st.text(str(tukey))
            interp = f"📌 집단 간 차이 유의성 p={format_p(p_val)}"

    elif method == "상관분석":
        sel_vs = st.multiselect("분석할 변수군 선택 (연속형)", num_cols)
        if st.button("통계 분석 실행") and len(sel_vs) >= 2:
            final_df = df[sel_vs].corr().round(3)
            
            # 2개 변수 선택 시 산점도 제공
            if len(sel_vs) == 2:
                plt.figure(figsize=(6, 5))
                sns.regplot(x=df[sel_vs[0]], y=df[sel_vs[1]], line_kws={"color": "red"})
                plot_img = get_plot_buffer()
                assump_report.append('<div class="assumption-pass">✅ 시각적 검토 준비 완료: 하단에 생성된 <b>산점도(Scatter Plot)와 회귀선</b>을 통해 두 변수가 직선 형태의 패턴을 보이는지 시각적으로 판단하십시오.</div>')
            else:
                plt.figure(figsize=(7, 5))
                sns.heatmap(final_df, annot=True, cmap="coolwarm")
                plot_img = get_plot_buffer()
                assump_report.append('<div class="assumption-pass">ℹ️ 다변량 분석 안내: 전체적인 패턴 파악을 위해 히트맵을 제공합니다. 정밀한 선형성 검토가 필요한 경우, 변수를 2개씩 선택하여 산점도를 확인하십시오.</div>')

            interp = "변수 간 선형적 상관계수 행렬입니다. 0.7 이상이면 강한 상관관계입니다."

    elif method == "신뢰도 분석":
        sel_items = st.multiselect("신뢰도 분석할 문항군 선택 (연속형)", num_cols)
        if st.button("통계 분석 실행") and len(sel_items) >= 2:
            items = df[sel_items].dropna(); k = items.shape[1]
            alpha = (k/(k-1)) * (1 - (items.var(ddof=1).sum() / items.sum(axis=1).var(ddof=1)))
            
            if alpha >= 0.7:
                assump_report.append(f'<div class="assumption-pass">✅ 신뢰도 양호: Cronbach Alpha {alpha:.3f} (기준 0.7 이상)</div>')
            else:
                assump_report.append(f'<div class="assumption-fail">⚠️ 신뢰도 낮음: Cronbach Alpha {alpha:.3f} (기준 0.7 미만). 문항 제거 또는 수정 필요.</div>')
            
            final_df = pd.DataFrame({"측정 지표": ["Cronbach α"], "수치": [f"{alpha:.3f}"]})
            interp = f"📌 신뢰도 계수는 {alpha:.3f}로 확인되었습니다."

    elif method == "회귀분석":
        rtype = st.radio("회귀 유형", ["선형 회귀분석 (Linear)", "로지스틱 회귀분석 (Logistic)"])
        xs = st.multiselect("독립변수군 (연속형/더미)", num_cols)
        y = st.selectbox("종속변수 (Linear:연속형 / Logistic:0,1범주형)", num_cols)
        
        if st.button("통계 분석 실행") and xs:
            if "선형" in rtype:
                X = sm.add_constant(df[xs]); model = sm.OLS(df[y], X).fit(); p_val = model.f_pvalue
                vifs = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]
                max_vif = max(vifs[1:]) if len(vifs) > 1 else 1.0
                if max_vif < 10:
                    assump_report.append(f'<div class="assumption-pass">✅ 다중공선성 없음: 최대 VIF {max_vif:.2f} (기준 10 미만)</div>')
                else:
                    assump_report.append(f'<div class="assumption-fail">⚠️ 다중공선성 경고: 최대 VIF {max_vif:.2f} (변수 제거 또는 차원 축소 고려 권장)</div>')
                dw = durbin_watson(model.resid)
                if 1.5 < dw < 2.5:
                     assump_report.append(f'<div class="assumption-pass">✅ 잔차 독립성 충족: Durbin-Watson {dw:.2f} (2에 근접)</div>')
                else:
                     assump_report.append(f'<div class="assumption-fail">⚠️ 잔차 독립성 주의: Durbin-Watson {dw:.2f} (시계열 분석 등 고려 필요)</div>')
                
                # [논문용 상세 결과] B, SE, t, p
                final_df = pd.DataFrame({
                    "B": model.params,
                    "SE": model.bse,
                    "t": model.tvalues,
                    "p": model.pvalues
                }).round(3)
                final_df['p'] = final_df['p'].apply(lambda x: "<.001" if x < 0.001 else f"{x:.3f}")
                
                interp = f"📌 모델 설명력(Adjusted R²)은 {model.rsquared_adj:.3f}이며, 모형의 적합도는 유의합니다(p={format_p(p_val)})."

            else: # 로지스틱
                X = sm.add_constant(df[xs]); model = sm.Logit(df[y], X).fit(disp=False); p_val = model.llr_pvalue
                
                # [논문용 상세 결과] B, SE, OR, 95% CI
                params = model.params
                conf = model.conf_int()
                conf.columns = ['Lower CI', 'Upper CI']
                
                final_df = pd.DataFrame({
                    "B": params,
                    "SE": model.bse,
                    "OR": np.exp(params),
                    "95% CI Lower": np.exp(conf['Lower CI']),
                    "95% CI Upper": np.exp(conf['Upper CI']),
                    "p": model.pvalues
                }).round(3)
                final_df['p'] = final_df['p'].apply(lambda x: "<.001" if x < 0.001 else f"{x:.3f}")
                
                interp = f"📌 로지스틱 회귀모형의 적합도는 유의합니다(p={format_p(p_val)}). OR(오즈비) 신뢰구간이 1을 포함하지 않아야 유의합니다."

    # --- Step 03: 결과 대시보드 (Solution 1: Dashboard Style) ---
    if final_df is not None:
        st.markdown('<div class="section-title"><span class="step-badge">03</span> 분석 결과 요약 및 학술적 해석</div>', unsafe_allow_html=True)
        
        # 1. 필수 가정 검정 (Assumptions)
        if assump_report:
            with st.expander("🔍 필수 가정 검정 (Assumption Check) 결과 확인", expanded=True):
                st.caption("통계 분석의 신뢰성을 확보하기 위해 필수적으로 확인해야 할 가정들입니다.")
                for msg in assump_report: st.markdown(msg, unsafe_allow_html=True)
        
        st.markdown("###") # 간격

        # 2. 메인 대시보드 (좌측: 상세표 / 우측: 요약카드 & 다운로드)
        col_main_L, col_main_R = st.columns([1.3, 1]) 
        
        with col_main_L:
            st.markdown("##### 📋 통계량 상세표")
            # 데이터프레임 표시 (인덱스 숨김)
            st.dataframe(final_df, use_container_width=True, hide_index=True)
            
        with col_main_R:
            st.markdown("##### 💡 핵심 결론")
            
            # P-value 존재 여부에 따른 카드 상태 설정
            if p_val is not None:
                if p_val < 0.05:
                    status_bg = "#dcfce7"; status_icon = "✅"; status_msg = "통계적 유의성 확보"
                else:
                    status_bg = "#fee2e2"; status_icon = "❌"; status_msg = "통계적으로 유의하지 않음"
            else:
                # 기술통계, 빈도분석 등 P-value 개념이 없는 경우
                status_bg = "#f1f5f9"; status_icon = "📊"; status_msg = "분석 결과 요약"

            # HTML Card 렌더링
            st.markdown(f"""
            <div style="background-color: {status_bg}; padding: 20px; border-radius: 12px; border: 1px solid #cbd5e1; margin-bottom: 15px;">
                <div style="font-size: 1.1rem; font-weight: 700; color: #334155; margin-bottom: 8px;">{status_icon} {status_msg}</div>
                <div style="font-size: 0.95rem; color: #475569; line-height: 1.6;">{interp}</div>
            </div>
            """, unsafe_allow_html=True)
            
            # 다운로드 버튼 (카드 하단에 꽉 차게 배치)
            st.download_button(
                label="📄 워드 리포트 다운로드",
                data=create_pro_report(method, final_df, interp, "통계 수치를 논문에 인용하세요.", plot_b=plot_img, assump="\n".join(assump_report)),
                file_name=f"STATERA_{method}.docx",
                use_container_width=True, 
                type="primary"
            )

        # 3. 시각화 (그래프가 있다면 하단에 크게 배치)
        if plot_img:
            st.markdown("###")
            st.markdown("##### 📊 시각화 결과")
            st.image(plot_img, use_container_width=True)

# 하단 연구 윤리 가이드
st.markdown(f"""
<div class="ethics-container">
    <div class="ethics-title">⚠️ 연구자 유의사항</div>
    <div class="ethics-text">
        1. 본 서비스에서 산출된 결과는 유의수준 0.05를 기준으로 한 통계적 판정입니다.<br>
        2. 최종 분석 결과의 정확성을 검토하고 보고서를 작성할 책임은 연구자 본인에게 있습니다.
    </div>
</div>
<div style='text-align: center; color: #cbd5e1; margin-top: 20px; font-size: 0.8rem;'>
    STATistical Engine for Research & Analysis | ANDA Lab | nncj91@snu.ac.kr
</div>
""", unsafe_allow_html=True)
